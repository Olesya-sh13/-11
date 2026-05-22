"""Конвертирует ЛР11_Отчёт_Качество_ПО.md в Word-документ."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

MD_FILE   = "ЛР11_Отчёт_Качество_ПО.md"
DOCX_FILE = "ЛР11_Отчёт_Качество_ПО.docx"

doc = Document()

# ── Настройки страницы (А4, поля 2 см) ──────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(section, attr, Cm(2))

# ── Стили шрифта ────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)

for h_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
    s = doc.styles[h_name]
    s.font.name = "Times New Roman"
    s.font.color.rgb = RGBColor(0, 0, 0)

doc.styles["Heading 1"].font.size = Pt(16)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 3"].font.size = Pt(13)
doc.styles["Heading 4"].font.size = Pt(12)

# ── Вспомогательные функции ──────────────────────────────────────

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p


def add_code_block(text):
    """Моноширинный блок для листингов."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    # Серый фон через shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(cell_text))
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = bold
    return row


def parse_inline(text):
    """Разбирает **bold** и *italic* внутри строки, возвращает список (text, bold, italic)."""
    parts = []
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        if m.group(1).startswith("**"):
            parts.append((m.group(2), True, False))
        elif m.group(1).startswith("*"):
            parts.append((m.group(3), False, True))
        else:
            parts.append((m.group(4), False, False))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts or [(text, False, False)]


def add_inline_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    for chunk, bold, italic in parse_inline(text):
        run = p.add_run(chunk)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p


# ── Основной парсер ─────────────────────────────────────────────

lines = open(MD_FILE, encoding="utf-8").readlines()

in_code   = False
code_buf  = []
in_table  = False
tbl_obj   = None
tbl_cols  = 0

i = 0
while i < len(lines):
    raw = lines[i].rstrip("\n")
    line = raw.strip()

    # ── Код-блок ────────────────────────────────────────────────
    if line.startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            add_code_block("\n".join(code_buf))
        i += 1
        continue

    if in_code:
        code_buf.append(raw)
        i += 1
        continue

    # ── Таблица ─────────────────────────────────────────────────
    if line.startswith("|"):
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Разделитель |---|---|
        if all(re.fullmatch(r":?-+:?", c) for c in cells if c):
            i += 1
            continue
        if not in_table:
            in_table = True
            tbl_cols = len(cells)
            tbl_obj  = doc.add_table(rows=0, cols=tbl_cols)
            tbl_obj.style = "Table Grid"
            add_table_row(tbl_obj, cells, bold=True)
        else:
            add_table_row(tbl_obj, cells)
        i += 1
        continue
    else:
        in_table = False
        tbl_obj  = None

    # ── Горизонтальная линия ─────────────────────────────────────
    if re.fullmatch(r"-{3,}|={3,}|\*{3,}", line):
        doc.add_paragraph()
        i += 1
        continue

    # ── Пустая строка ────────────────────────────────────────────
    if not line:
        i += 1
        continue

    # ── Заголовки ────────────────────────────────────────────────
    m = re.match(r"^(#{1,4})\s+(.*)", line)
    if m:
        level = len(m.group(1))
        add_heading(m.group(2), level=level)
        i += 1
        continue

    # ── Маркированный список ─────────────────────────────────────
    m = re.match(r"^[-*•]\s+(.*)", line)
    if m:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for chunk, bold, italic in parse_inline(m.group(1)):
            run = p.add_run(chunk)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold   = bold
            run.italic = italic
        i += 1
        continue

    # ── Нумерованный список ──────────────────────────────────────
    m = re.match(r"^\d+[.)]\s+(.*)", line)
    if m:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        for chunk, bold, italic in parse_inline(m.group(1)):
            run = p.add_run(chunk)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold   = bold
            run.italic = italic
        i += 1
        continue

    # ── Обычный абзац с inline-разметкой ─────────────────────────
    add_inline_paragraph(line)
    i += 1

doc.save(DOCX_FILE)
print(f"Сохранено: {DOCX_FILE}")
